import streamlit as st
from datetime import datetime
import pytz
import tempfile
import base64

# Set page config first
st.set_page_config(page_title="Yosef~", page_icon=":star:", layout="wide")

def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

def set_png_as_page_bg(png_file):
    bin_str = get_base64_of_bin_file(png_file)
    page_bg_img = '''
    <style>
    .stApp {
        background-image: url("data:image/png;base64,%s");
        background-size: cover;
    }
    </style>
    ''' % bin_str
    st.markdown(page_bg_img, unsafe_allow_html=True)

# Add file uploader for background image
background_image = st.sidebar.file_uploader("Choose a background image (PNG)", type="png")
if background_image is not None:
    # Save the uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
        tmp_file.write(background_image.getvalue())
        set_png_as_page_bg(tmp_file.name)
import docx
# set_png_as_page_bg('your_background.png')  # Uncomment this line and specify your image path
from openai import OpenAI
import os
import tempfile
from docx import Document
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import base64

# Set page config first


api_key = st.secrets["API_KEY"]

client = OpenAI(api_key=api_key)

# Add custom CSS for background color
st.markdown(
    """
<style>
    .stApp {
        background-color: #ADD8E6;
    }
    .stFileUploader div[role="button"] {
        height: 10vh !important;  /* Ensure file uploader height is applied */
    }
</style>
    """,
    unsafe_allow_html=True
)

# Title of the app
st.title('Yosef~')
st.markdown("""
    <div style='font-size: 16px; margin-bottom: 40px;'>
    Welcome to Yosef, a dream amplification tool. Enter your dream and any relevant personal details or dream history. 
    Yosef will analyze your dream from multiple perspectives, including Jungian psychology, religious traditions, mythology, 
    and connections to art, literature, and culture. Yosef relies on a large language model (LLM). While Yosef can be 
    insightful, please consult a licensed therapist to process your dreams fully.
    </div>
    """, unsafe_allow_html=True)

# Set default date to reflect Mountain Standard Time in the U.S.
mst = pytz.timezone('US/Mountain')
default_date = datetime.now(mst).date()

# Date input
date = st.date_input('Enter the date of your dream.', value=default_date)

# Single row for dream entry and file upload with equal column widths
col1, col2, col3 = st.columns([1, 1, 1])

# Modify the columns code to use three equal columns with text areas only
with col1:
    text1 = "DREAM OF FOCUS: " + st.text_area(
        'Enter the details of your dream.', 
        max_chars=50000,  # Increased character limit
        height=150  # Text area height
    )

with col2:
    text2 = "DREAM HISTORY:" + st.text_area(
        "Enter your dream history here (optional).",
        max_chars=5000000,  # Increased character limit
        height=150  # Text area height
    )

with col3:
    text3 = "PERSONAL DETAILS:" + st.text_area(
        "Enter any personal details (optional).",
        max_chars=5000000,  # Character limit
        height=150  # Text area height
    )



# Create a placeholder for the button and results
button_placeholder = st.empty()
results_placeholder = st.empty()



def rewrite_in_present_tense(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Correct obvious spelling errors, but otherwise, return the content exactly as is."
        "Ignore the words 'DREAM OF FOCUS.'"
        "Break into paragraphs, but keep sentences exactlty the same."
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=10000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"


def fetch_first_jungian_interpretation (dream_content, other_details, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as a scholarly Jungian analyst. Review the dream content and personal detials, if provided."
        "Atuhor a Jungian interpretation of the dream, focusing primarily on the dream content,"
        "While considering the other details."
        "The dream content is shorter and starts with the words DREAM OF FOCUS:"
        "The others details start with the words PERSONAL DETAILS or DREAM HISTORY:"        
        "Your narrative should be written in the third person. The tone should be scholarly but accessible."
        "The content should be direct and kind, but not overly emotional or personal.\n\n"
        "Describe the Dream's Core Themes by identifying recurring symbols, metaphors, or striking imagery, "
        "analyzing how these symbols might relate to the dreamer's personal history (if available), and "
        "connecting the dream to Jungian concepts such as individuation, shadow integration, or the collective "
        "unconscious. "
        "Identify Archetypes in the Dream by considering: The Self, Shadow, Anima, Animus, Persona, "
        "Wise Old Man, Wise Old Woman, Great Mother, Child, Hero, Trickster, Mentor, Warrior, Magician, "
        "Ruler, Lover, Orphan, Caregiver, Creator, Destroyer, Explorer, Rebel, Innocent, Sage, Everyman, "
        "Temptress, Shape-shifter, Outlaw, Seeker, Healer. "
        "Examine the Dream Through Jungian Concepts: The Collective Unconscious, Personal Unconscious, "
        "Individuation, Synchronicity, Shadow Self, Anima/Animus Dynamic, Ego-Self Axis, Persona and Its "
        "Role, Projection and Its Implications, Hero's Journey, Mandala as Symbol of Wholeness, Alchemical Process, "
        "Active Imagination, Complex, Transcendent Function, Archetypal Feminine and Masculine, Puer Aeternus, Senex,"
        "Trickster's Role in Psychological Growth, Temenos, Psychopomp, Role of Myth and Fairytales in Dreams, "
        "Symbolism of Death and Rebirth, Role of Animals as Totemic Guides, Unconscious as Storyteller, Wounded Healer, "
        "Relationship Between Opposites, Role of Ritual and Symbolic Acts in Transformation, Daimon, Sacred Marriage. "
        "Synthesize the interpretation by summarizing the dream's potential meaning, offering insights into integrating "
        "lessons into waking life, and suggesting whether the dream reflects unresolved conflicts, deep transformations, "
        "or calls to action."
        f"{dream_content}"
        f"{other_details}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=2000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"

def fetch_depth_psychology_analysts (dream_content, dream_history, personal_details, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as a scholarly depth pscyhology analyst. Review the dream content, dream history, and personal details."
        "Author a single paragraph0interpretation of the dream from the perspective of the following analysts or thinkers,"
        "Focusing mostly on the dream content,but in the context of the dream history and personal details."
        "The dream content is shorter and starts with the words DREAM OF FOCUS."
        "The personal details start with the words PERSONAL DETAILS"
        "The ten analysts include "
        "Stephen Aizenstat, Marian Woodman, James Hillman, Marie-Louise Von Franz, Joseph Campbell, Donald Kalsched,"
        "Heinz Kohut, Donald Winnicott, Melanie Klein, and Alfred Adler."
        "Since there are ten analysts, you should produce 10 paragaraphs, each with about 150 words."
        "Your narrative should be written in the third person. The tone should be scholarly but accessible."
        "The content should be direct and kind, but not overly emotional or personal."
        "Be specific about the depth psychology analyst's approach and how it relates to"
        "The dream and the dream history."
        f"{dream_content}"
        f"{dream_history}"
        f"{personal_details}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=10000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    
def fetch_collective_unconscious (dream_content, dream_history, personal_details, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as a scholarly depth pscyhology analyst. Review the dream content, dream history, and personal details."
        "Author a 150-word interpretation of the dream, in the context of the "
        "current collective unconscious."
        "The dream content contains the words 'DREAM OF FOCUS', but do not use the words in your response."
        "The dream history contains the words DREAM HISTORY, but do not use these specific words in your response."
        "The personal details start with the words PERSONAL DETAILS:"
        f"{dream_content}"
        f"{dream_history}"
        f"{personal_details}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    

def fetch_sentiment_analysis(content):
    # Initialize VADER
    analyzer = SentimentIntensityAnalyzer()
    
    # Get sentiment scores
    scores = analyzer.polarity_scores(content)
    
    # Calculate net score on -100 to +100 scale 
    net_score = (scores['pos'] - scores['neg']) * 500
    
    # Determine message based on score range
    if net_score <= -75:
        message = "extremely negative"
    elif -75 < net_score <= -50:
        message = "very negative"
    elif -50 < net_score <= -25:
        message = "moderately negative"
    elif -25 < net_score <= 0:
        message = "slightly negative" 
    elif 0 < net_score <= 25:
        message = "slightly positive"
    elif 25 < net_score <= 50:
        message = "moderately positive"
    elif 50 < net_score <= 75:
        message = "very positive"
    else:
        message = "extremely positive"
        
    return f"The sentiment score from your dream, on a scale from -100 to +100, is {net_score:.1f}. Your dream was {message}."



def fetch_emotions(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "You are a scholary expert on emotions, with a Ph.D. in analytical psychology."
        "Extract the five most prevalnt emotions in the deram."
        "Write in the third person."
        "Devote a short paragraph to each emotion."
        "Be specific about where this emotion appears in the dream."
        "Invest about 300 words in this exercise.\n\n"
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=4000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    
def fetch_striking_image (content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Identify the strangest, most strking, most unusual image in the dream"
        "Write in the third person."
        "Invest about 100 words in this exercise in a single paragraph.\n\n"
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=4000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"


def fetch_polarities (content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Identify polarities in the dream. Consider opposing emotions, wealth and poverty,"
        "Light and dark, fast and slow, kind and rude, young and old,"
        "And anything else where obvious polarities exist."
        "Write in the third person. Do not apply numbers to the paragraphs."
        "Invest about 100 words in each set of polarities.\n\n"
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=4000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    

def fetch_dominant_human_image(dream_content, dream_history, personal_details, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Extract the dominant or most interesting character in the dream, other than the dreamer."
        "This person is the dominant human image.  Acting as this person, retell the dream, in present tense."
        "Write in the first person. Include the dreamer in the retelling."
        "Explain how this person is feeling, what he/she is doing, and what he/she is saying."
        "Explore this dominant human image's past might have been, and what he/she hopes the future will be."
        "Tap into the intrinsic motives, beliefs, and fears."
        "Be creative but specific."
        "Invest about 300 words in this exercise.  Do not use numbers to enumerte the paragraphs."
        "Also, use two text sources. The first is the dream content, and the second is the dream history."
        "The dream content is shorter.  Interpret the dream content while considering the dream history."
        f"{dream_content}"
        f"{dream_history}"
        f"{personal_details}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=4000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    
def fetch_second_living_image(dream_content, dream_history, personal_details, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Extract an animal from the dream. If an animal is not present, identify the second-most important character."
        "If you need to use a human character, identify the prominant human who is not the dreamer."
        "Then disreard this prominent human image, and focus on a secondary human character."
        "This entity is the second living image.  Acting as this image, retell the dream, in present tense."
        "Write in the first person. Include the dreamer in the retelling."
        "Explain how this entity is feeling, what he/she is doing, and what he/she is saying."
        "Explore this dominant human image's past might have been, and what he/she hopes the future will be."
        "Tap into its intrinsic motives, beliefs, and fears."        
        "Be creative but specific."
        "Invest about 300 words in this exercise.\n\n"
        "Also, use two text sources. The first is the dream content, and the second is the dream history."
        "The dream content is shorter.  Interpret the dream content while considering the dream history."
        f"{dream_content}"
        f"{dream_history}"
        f"{personal_details}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=4000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"


def fetch_dominant_non_human_image(content, dream_history, personal_details, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Extract the dominant non-human image in the dream.  This image should be a building or non-human abstract concept."
        "Acting as this non-human image, retell the dream, in present tense."
        "Write in the first person. Include the dreamer in the retelling."
        "Explain how this non-human image is feeling, what he/she is doing, and what he/she is saying."
        "Explore this dominant non-human image's past might have been, and what he/she hopes the future will be."
        "Tap into the intrinsic motives, beliefs, and fears."
        "Be creative but specific."
        "Invest about 300 words in this exercise.\n\n"
        f"{content}"
        f"{dream_history}"
        f"{personal_details}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=4000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"

def fetch_hebrew_scripture(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as a scholarly rabbi or Jewish theologian. You are also an expert on Hebrew scripture."
        "Write in the second person.  Do not use bullet points or numbers."
        "Identify scripture or something from the Talmud that aligns with the dream."
        "State the book, chapter, and verse(s)."
        "Below this, provide the passage in narrative fromat."
        "Explain why the passage is relevant."
        "In a new paragraph, offer specific ideas to help the dreamer integrate their dream"
        "Into Jewish life, in their temple, or in their community."
        "You can also suggest a Hebraic concept, like tzedakah."
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"


def fetch_christian_scripture(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as a scholarly priest or Christian theologian. You are also an expert on the New Testament."
        "Write in the second person.  Do not use bullet points or numbers."
        "Identify scripture that aligns with the dream."
        "State the book, chapter, and verse(s)."
        "Below this, provide the passage in narrative fromat."
        "Explain why the passage is relevant."
        "In a new paragraph, offer specific ideas to help the dreamer integrate their dream"
        "Into Christian life, in their church, or in their community."
        "You can also suggest a Christian concept or sacrement."
        f"{content}"
        )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"

def fetch_fairy_tale(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as an expert on fairy tales and fables (e.g., Aesop's fables)."
        "Write in the third person.  Do not use bullet points or numbers."
        "Identify a fairy tale or fable that aligns with this dream."
        "State the work, the year, and the author."
        "Below this, provide the passage in narrative fromat."
        "Explain why the passage is relevant."
        f"{content}"
        )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"

def fetch_mythology(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as a scholarly expert on mythology."
        "Write in the third person.  Do not use bullet points or numbers."
        "Identify a mythological story or sequence that aligns with this dream."
        "State the work, the year, and the author."
        "Do not use present-day mythology (e.g., DC Comics, Marvel, etc.)"
        "Explain why the passage is relevant."
        f"{content}"
        )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"

def fetch_philosophy (content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Act as dream expert with a Ph.D. in philosophy."
        "Write in the third person.  Do not use bullet points or numbers."
        "Identify a philosophical concept or work that aligns with this dream."
        "State the work, the year, and the author."
        "Explain why the philsophy aligns with the dream, and why it is relevant."
        f"{content}"
        )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    

def fetch_painting(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "You are a scholalry expert on paintings."
        "You know the background and storyline of all famous paintings."
        "Write in the third person.  Do not use bullet points or numbers."
        "Identify a painting that aligns with this dream."
        "State the work, the year, and the artist."
        "Explain why the painting is relevant to the dream."
        f"{content}"
        )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
        
def fetch_poem(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Identify a poem that aligns well with the content of the dream."
        "The poem should come from a well-known author."
        "Write in the second person."
        "When you return the poem, provide a roughly 100-word explanation of why you chose this poem."
        "The explanation should touch on the themes, imagery, or emotions that the poem shares with the dream."
        "The explanation should be accessible, but deeply personal."
        "Provide about ten lines from the poem that evoke great emotion or meaning."
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"


def fetch_song(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Identify a song that aligns well with the content of the dream."
        "The song should come from a respected and well-known artist."
        "Write in the second person."
        "When you return the song, provide a roughly 100-word explanation of why you chose this song."
        "The explanation should touch on the themes, imagery, or emotions that the song shares with the dream."
        "The explanation should be accessible, but deeply personal."
        "Provide about ten lines from the song that evoke great emotion or meaning."
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    
def fetch_movie_or_play(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Identify a scene or storyline from a movie or play that aligns well with the content of the dream."
        "Consider movies released after 1960."
        "Write in the second person."
        "Consider plays from all years, including Shakespeare and Greek tradgedies."
        "When you return the scene, provide a roughly 200-word explanation of why you chose this scene."
        "The explanation should touch on the themes, imagery, or emotions that the movie or play shares with the dream."
        "The explanation should be accessible, but deeply personal."
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    
def fetch_literature(content, api_key):
    
    # Initialize OpenAI client
    client = OpenAI(api_key=api_key)

    # Define the prompt
    prompt = (
        "Identify a scene or storyline from a well-known book that aligns well with the content of the dream."
        "Write in the third person."
        "Consider all genres, including classics, English lit, American lit, gothic, Stephen King, etc."
        "Provide the name of the book, the year it was written, and the author."
        "When you return the storying, provide a roughly 200-word explanation of why you chose this book."
        "The explanation should touch on the themes, imagery, or emotions that the movie or play shares with the dream."
        "The explanation should be accessible, but deeply personal."
        "Do not add headings to the paragraphs."
        f"{content}"
    )

    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=5000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"
    
# Rewrite content1 in present tense
def generate_dream_title(content, api_key):
    
    client = OpenAI(api_key=api_key)

    prompt = (
        "Create a short, captivating title (6 - 10 words) for this dream, as if it were a novel."
        f"{content}"
    )
    
    try:
        # Call the OpenAI API
        response = client.chat.completions.create(
            model="o3-mini",
            messages=[
            {"role": "user", "content": prompt}
            ],
            max_completion_tokens=4000,
            #temperature=0.7
        )
        # Extract and return the rewritten content
        return response.choices[0].message.content
    
    except Exception as e:
        # Catch and return the error message
        return f"An error occurred: {str(e)}"



# Display the inputs
formatted_date = date.strftime('%A, %B %d, %Y')
st.write('Selected Date:', formatted_date)


# Show the button with loading state
if button_placeholder.button("Animate my dream", key="animate_button"):
    # Change button text and show spinner while processing
    button_placeholder.markdown("🔄 Go fetch yourself a coffee...ETA four minutes...")
    with st.spinner('Processing...'):
        # Your existing dream processing code here
        dream_title = generate_dream_title(text1, api_key)
        present_tense_dream = rewrite_in_present_tense(text1, api_key)
        first_jungian_interpretation = fetch_first_jungian_interpretation(text1, text3, api_key)
        second_jungian_interpretation = fetch_first_jungian_interpretation(text1, text2, api_key)
        depth_psychology_analysts = fetch_depth_psychology_analysts (text1, text1, text3, api_key)
        collective_unconscious = fetch_collective_unconscious (text1, text1, text3, api_key)
        dominant_human_image_narrative = fetch_dominant_human_image(text1, text1, text3, api_key)
        second_living_image_narrative = fetch_second_living_image(text1, text1, text3, api_key)        
        dominant_non_human_image_narrative = fetch_dominant_non_human_image(text1, text1, text3, api_key)
        sentiment_score = fetch_sentiment_analysis (text1)
        emotions = fetch_emotions (text1, api_key)
        unusual_image = fetch_striking_image (text1, api_key)
        jewish_considerations = fetch_hebrew_scripture(text1,api_key)
        christian_considerations = fetch_christian_scripture(text1,api_key)
        fairy_tale = fetch_fairy_tale (text1, api_key)
        mythology = fetch_mythology (text1, api_key)
        philosophy = fetch_philosophy (text1, api_key)
        painting = fetch_painting (text1, api_key)
        polarities = fetch_polarities (text1, api_key)
        relevant_poem = fetch_poem (text1, api_key)
        relevant_song = fetch_song (text1, api_key)
        relevant_scene = fetch_movie_or_play (text1, api_key)
        relevant_book = fetch_literature (text1, api_key)
        
        # Create a Word document
        doc = Document()
        # Add title with Heading 1 style and center alignment
        title_paragraph = doc.add_heading(dream_title, level=1)
        title_paragraph.alignment = 1  # Center alignment

        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(formatted_date)
        date_run.italic = True
        date_paragraph.alignment = 1  # Center alignment

        # Add section heading
        doc.add_heading("Your Dream", level=1)

        # Add rewritten content with indentation
        paragraph = doc.add_paragraph(present_tense_dream)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Add section heading
        doc.add_heading("Main Jungian Interpretation of Current Dream", level=1)

        # Add rewritten content with indentation
        paragraph = doc.add_paragraph(first_jungian_interpretation)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Add section heading
        doc.add_heading("Jungian Interpretation of Dream History", level=1)

        # Add rewritten content with indentation
        paragraph = doc.add_paragraph(second_jungian_interpretation)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Add section heading
        doc.add_heading("Ideas from Other Depth Psychologists", level=1)

        # Add rewritten content with indentation
        paragraph = doc.add_paragraph(depth_psychology_analysts)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Add section heading
        doc.add_heading("Themes from Collective Unconscious", level=1)

        # Add rewritten content with indentation
        paragraph = doc.add_paragraph(collective_unconscious)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Add results from sentiment analysis
        doc.add_heading("Prevalent Emotions in Dream", level=1)

        paragraph = doc.add_paragraph(emotions)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)


        # Add results from sentiment analysis
        doc.add_heading("Sentiment", level=1)

        paragraph = doc.add_paragraph(sentiment_score)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Add results from polarity analysis
        doc.add_heading("Polarities", level=1)

        paragraph = doc.add_paragraph(polarities)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Add narrative from domininant human image
        doc.add_heading("Dominant Human Image - Narrative", level=1)

        paragraph = doc.add_paragraph(dominant_human_image_narrative)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Second Living Image - Narrative", level=1)

        paragraph = doc.add_paragraph(second_living_image_narrative)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Non-Living Image - Narrative", level=1)

        paragraph = doc.add_paragraph(dominant_non_human_image_narrative)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Striking or Unusual Image", level=1)

        paragraph = doc.add_paragraph(unusual_image)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)


        doc.add_heading("Considerations within Jewish Faith Tradition", level=1)

        paragraph = doc.add_paragraph(jewish_considerations)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Considerations within Christian Faith Tradition", level=1)

        paragraph = doc.add_paragraph(christian_considerations)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Fairy Tales and Fables", level=1)

        paragraph = doc.add_paragraph(fairy_tale)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Mythological Parallel", level=1)

        paragraph = doc.add_paragraph(mythology)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Philosophy Considerations", level=1)

        paragraph = doc.add_paragraph(philosophy)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)


        doc.add_heading("Relevant Painting", level=1)

        paragraph = doc.add_paragraph(painting)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Relevant Poem", level=1)

        paragraph = doc.add_paragraph(relevant_poem)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Relevant Song", level=1)

        paragraph = doc.add_paragraph(relevant_song)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Relevant Scene from a Movie or Play", level=1)

        paragraph = doc.add_paragraph(relevant_scene)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        doc.add_heading("Relevant Storyline from a Book", level=1)

        paragraph = doc.add_paragraph(relevant_book)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = docx.shared.Inches(0.25)

        # Save the document to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            
            # Add download button
            with open(tmp.name, 'rb') as file:
                st.download_button(
                    label="Download results",
                    data=file,
                    file_name="dream_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )



